from docx import Document
from docx.shared import Inches

from docx.shared import Pt

document = Document()

style = document.styles['Normal']
font = style.font

font.name = 'Arial'
font.size = Pt(9)


document.add_heading('Plantillas de Configuración Básica de Gestión para Equipamiento Cisco', 0)
document.add_page_break()

document.add_heading('PLANTILLAS DE CONFIGURACIÓN', 1)
document.add_heading('Configuración común a todos los servicios', 2)

document.add_heading('Definición del banner', 3)
document.add_paragraph('Desde el modo de configuración del equipo, se especifica el banner que se mostrará antes de que el usuario introduzca sus credenciales de acceso: ')
document.add_paragraph('banner login @ {{ Texto_Banner_Pre_Autenticacion }} @ ')
document.add_paragraph('A continuación, se introduce el banner que se mostrará una vez que el usuario haya sido autenticado correctamente: ')
document.add_paragraph('banner exec @ {{ Texto_Banner_Post_Autenticacion }} @ ')
document.add_paragraph('NOTA: La escritura de los <Texto_Banner> debe comenzar y terminar con un <Caracter_Delimitador> (%, &, $, @, …) para funcionar correctamente. Además, este carácter no se puede utilizar dentro de <Texto_Banner>. Durante la introducción del texto es posible utilizar la tecla Enter para producir saltos de línea. Cada línea admite un máximo de 255 caracteres. ')

document.add_heading('Configuración de fecha y zona horaria', 3)
document.add_paragraph('Desde el modo enable, se ajusta la fecha y hora locales del equipo. ')
document.add_paragraph('clock set {{ hora_minuto_segundo }} {{ dia }} {{ mes }} {{ ano }}')
document.add_paragraph('Desde el modo de configuración, se configura la zona horaria y los parámetros del cambio de horario de verano. ')
document.add_paragraph(' clock timezone MET 1 \n clock summer-time METDST recurring last Sun Mar 2:00 last Sun oct 3:00')

document.add_heading('Nombre del equipo', 3)
document.add_paragraph('Asignación del mnemónico del EDC. Desde el modo de configuración:  ')
document.add_paragraph('hostname {{ Mnemonico_EDC }} ')

document.add_heading('Desactivación de servicios no necesarios', 3)
document.add_paragraph('Por motivos de seguridad, se desactivarán una serie de protocolos y servicios que no van a ser utilizados en el EDC. Desde el modo de configuración: ')
document.add_paragraph(' no cdp run \n no service tcp-small-servers \n no service udp-small-servers no service finger \n no service dhcp \n {% if modelo_equipo != "Cisco C1000"  %} \n no ip source-route \n {% endif %} \n no ip source-route \n no ip domain-lookup \n no service config \n no ip bootp server   !!! * \n no ip name-server ')
document.add_paragraph('* Determinados modelos no admiten este comando al no disponer de la funcionalidad, por lo que no es problemático si al introducirlo da error. ')
document.add_paragraph('Configuración de otros parámetros generales. Desde el modo de configuración: ')
document.add_paragraph(' service nagle \n service tcp-keepalives-in \n ip subnet-zero \n service timestamps debug datetime localtime msec show-timezone \n service timestamps log datetime localtime msec show-timezone')

document.add_heading('Gestión por loopback', 3)
document.add_paragraph('Se crea la interfaz loopback de gestión.  ')
document.add_paragraph(' interface loopback 600 \n description Direccion IP de gestion EDC \n ip address {{IPGestion}} 255.255.255.255 \n no ip directed-broadcast \n no ip proxy-arp \n no shutdown \n exit ')
document.add_paragraph(' ip tftp source-interface Loopback 600 \n ip ftp source-interface Loopback 600 \n ip tacacs source-interface Loopback 600 \n logging source-interface Loopback 600 \n snmp-server trap-source Loopback 600 \n ntp source Loopback 600 ')
document.add_paragraph('Donde: <IPGestión> Es la dirección IP utilizada para la gestión del EDC. Por tanto, debe ser única en todo el ámbito del servicio.')



document.add_page_break()

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run('and some ')
p.add_run('italic.').italic = True

p = document.add_paragraph('')

document.add_heading('Heading, level 1', level=1)
document.add_heading('Heading, level 1', level=2)
document.add_heading('Heading, level 1', level=3)
document.add_heading('Heading, level 1', level=4)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph('first item in unordered list', style='List Bullet')

document.add_paragraph('first item in ordered list', style='List Number')
document.add_paragraph('second item in ordered list', style='List Number')
document.add_paragraph('second item in ordered list', style='List Number')






document.add_page_break()

document.save('demo.docx')