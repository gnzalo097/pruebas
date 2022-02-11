def limpiar_linea(linea, etiqueta):
    new_line1 = linea.replace('{# ', '')
    new_line2 = new_line1.replace(' -#}', '')
    new_line3 = new_line2.replace('\n', '')
    new_line4 = new_line3.replace(etiqueta, '')
    return new_line4

def limpiar_saltos(linea):
    new_line1 = linea.replace('\n', '')
    return new_line1


# file_output = open(r"C:\Users\Gonzalo\Desktop\prueba_py\probando_py\word_automatico.py", "w")
file = open("template.j2", "r")

# lines = file.readlines()
# print(lines)


from docx import Document
from docx.shared import Inches
from docx.shared import Pt

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(9)



contador_comandos = 0


array_comandos = []

contador_apartado = 0
contador_sub_apartado = 0
contador_sub_sub_apartado = 0
contador_sub_sub_sub_apartado = 0


for line in file:

    if "[titulo]" in line:
        print("[titulo]")

        new_line = limpiar_linea(line, "[titulo] ")

        
        document.add_heading( new_line , 0)
        document.add_page_break()
        

    elif "[apartado]" in line:
        print("[apartado]")

        contador_apartado += 1

        new_line = limpiar_linea(line, "[apartado] ")
        document.add_heading(str(contador_apartado) + ". " + new_line , 1)

        
    elif "[sub-apartado]" in line:
        print("[sub-apartado]")

        contador_sub_apartado += 1

        new_line = limpiar_linea(line, "[sub-apartado] ")
        document.add_heading(str(contador_apartado) + "." + str(contador_sub_apartado) + ". " + new_line , 2)


              
    elif "[sub-sub-apartado]" in line:
        print("[sub-sub-apartado]")

        contador_sub_sub_apartado += 1

        new_line = limpiar_linea(line, "[sub-sub-apartado] ")
        document.add_heading(   str(contador_apartado) +
                                "." +
                                str(contador_sub_apartado) +
                                "." + 
                                str(contador_sub_sub_apartado) + 
                                ". " + 
                                new_line, 3)


        
    elif "[sub-sub-sub-apartado]" in line:
        print("[sub-sub-sub-apartado]")

        contador_sub_sub_sub_apartado += 1

        new_line = limpiar_linea(line, "[sub-sub-sub-apartado] ")
        document.add_heading( str(contador_apartado) + "." + str(contador_sub_apartado) + "." + str(contador_sub_sub_apartado) + "." + str(contador_sub_sub_sub_apartado) + ". " + new_line , 4)   

                
    elif "[info]" in line:
        print("[info]")

        new_line = limpiar_linea(line, "[info] ")

        document.add_paragraph( new_line )

    elif line == "\n" or line == " ":
        print("vacio")

    else:
        print("---- COMANDO")

        # contador_comandos += 1

        new_line = limpiar_saltos(line)

        document.add_paragraph( new_line )


document.save('word_automatico2_render.docx')

file.close() 





    