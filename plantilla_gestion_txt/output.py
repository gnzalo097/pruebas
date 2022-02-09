from docx import Document
from docx.shared import Inches
from docx.shared import Pt

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(9)


document.add_heading('Plantillas de Configuración Básica de Gestión para Equipamiento Cisco', 0)
document.add_page_break()

document.add_heading('1.  PLANTILLAS DE CONFIGURACION', 1)

document.add_heading('1.1. 	Configuración común a todos los servicios', 2)

document.add_heading('1.1.1.  Definición del banner', 3)
document.add_paragraph('Desde el modo de configuración del equipo, se especifica el banner que se mostrará antes de que el usuario introduzca sus credenciales de acceso:')
document.add_paragraph('banner login @ {{ Texto_Banner_Pre_Autenticacion }} @ ')
document.add_paragraph('A continuación, se introduce el banner que se mostrará una vez que el usuario haya sido autenticado correctamente:')
document.add_paragraph('banner exec @ {{ Texto_Banner_Post_Autenticacion }} @ ')
document.add_paragraph('NOTA: La escritura de los <Texto_Banner> debe comenzar y terminar con un <Caracter_Delimitador> (%, &, $, @, …) para funcionar correctamente. Además, este carácter no se puede utilizar dentro de <Texto_Banner>. Durante la introducción del texto es posible utilizar la tecla Enter para producir saltos de línea. Cada línea admite un máximo de 255 caracteres.')

document.add_heading('1.1.2.  Configuración de fecha y zona horaria', 3)
document.add_paragraph('Desde el modo enable, se ajusta la fecha y hora locales del equipo.')
document.add_paragraph('clock set {{ hora_minuto_segundo }} {{ dia }} {{ mes }} {{ ano }}')
document.add_paragraph('Desde el modo de configuración, se configura la zona horaria y los parámetros del cambio de horario de verano.')
document.add_paragraph('{% if localizacion_canarias == false %}')
document.add_paragraph('clock timezone MET 1')
document.add_paragraph('{% endif %}')
document.add_paragraph('clock summer-time METDST recurring last Sun Mar 2:00 last Sun oct 3:00 ')

document.add_heading('1.1.3. Nombre del equipo', 3)
document.add_paragraph('hostname {{ Mnemonico_EDC }} ')

document.add_heading('1.1.4.  Desactivacion de servicios no necesarios', 3)
document.add_paragraph('Por motivos de seguridad, se desactivarán una serie de protocolos y servicios que no van a ser utilizados en el EDC. Desde el modo de configuración:')
document.add_paragraph('no cdp run')
document.add_paragraph('no service tcp-small-servers ')
document.add_paragraph('no service udp-small-servers ')
document.add_paragraph('no service finger')
document.add_paragraph('no service dhcp')
document.add_paragraph('{% if modelo_equipo != "Cisco C1000"  %}')
document.add_paragraph('no ip source-route           ')
document.add_paragraph('{% endif %}')
document.add_paragraph('no ip domain-lookup')
document.add_paragraph('no service config')
document.add_paragraph('no ip bootp server   !!! * ')
document.add_paragraph('no ip name-server')
document.add_paragraph('* Determinados modelos no admiten este comando al no disponer de la funcionalidad, por lo que no es problemático si al introducirlo da error.')
document.add_paragraph('Configuración de otros parámetros generales. Desde el modo de configuración:')
document.add_paragraph('service nagle')
document.add_paragraph('service tcp-keepalives-in')
document.add_paragraph('ip subnet-zero')
document.add_paragraph('service timestamps debug datetime localtime msec show-timezone')
document.add_paragraph('service timestamps log datetime localtime msec show-timezone')

document.add_heading('1.1.5.  Gestión por loopback', 3)
document.add_paragraph('Se crea la interfaz loopback de gestión.')
document.add_paragraph('interface loopback 600')
document.add_paragraph('    description Direccion IP de gestion EDC')
document.add_paragraph('    ip address {{IPGestion}} 255.255.255.255')
document.add_paragraph('    no ip directed-broadcast')
document.add_paragraph('    no ip proxy-arp')
document.add_paragraph('    no shutdown')
document.add_paragraph('exit')
document.add_paragraph('ip tftp source-interface Loopback 600')
document.add_paragraph('ip ftp source-interface Loopback 600')
document.add_paragraph('ip tacacs source-interface Loopback 600')
document.add_paragraph('logging source-interface Loopback 600')
document.add_paragraph('snmp-server trap-source Loopback 600')
document.add_paragraph('ntp source Loopback 600')
document.add_paragraph('Donde: <IPGestión> Es la dirección IP utilizada para la gestión del EDC. Por tanto, debe ser única en todo el ámbito del servicio.')

document.save('demo_2.docx')
